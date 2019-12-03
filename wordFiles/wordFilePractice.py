import docx
import os

os.chdir('D:\Dropbox\CodingResources\webDevelopmentProjects\Mine\practice-python-exercises\wordFiles')

d = docx.Document('D:\Dropbox\CodingResources\webDevelopmentProjects\Mine\practice-python-exercises\wordFiles\demo.docx')

print(d.paragraphs)
print(d.paragraphs[0])
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

#runs for each change in font style
p = d.paragraphs[1]
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)
# check if font style is present
print(p.runs[1].bold)
print(p.runs[3].italic)

#change font style 
p.runs[3].underline = True
#change text in file
p.runs[3].text = 'italic and underline'


#make changes and save to separate file
d.save('demo2.docx')

p.style = 'Title'
d.save('demo3.docx')

#write new paragraphs and save to file
d = docx.Document()
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph.')

d.save('demo4.docx')

p = d.paragraphs[0]
p.add_run('This is a new run.')

p.runs

p.runs[1].bold = True
d.save('demo5.docx')